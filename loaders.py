from langchain.schema import Document as LangchainDocument
from PyPDF2 import PdfReader
from docx import Document
from dspy_pipeline import IEPPipeline
import os
from typing import List, Optional
import logging
import re

logger = logging.getLogger(__name__)

def process_with_dspy(documents: List[LangchainDocument]) -> List[LangchainDocument]:
    """Process documents using the DSPy pipeline.
    
    Args:
        documents (List[LangchainDocument]): List of documents to process
        
    Returns:
        List[LangchainDocument]: Processed documents with enhanced content
    """
    iep_pipeline = IEPPipeline()
    results = []

    for doc in documents:
        try:
            # Process with DSPy
            iep_result = iep_pipeline.forward(doc.page_content)
            
            # Create enhanced content
            enhanced_content = f"""
Original Content:
{doc.page_content}

Key Insights:
{iep_result['insights']}

Important Entities:
{iep_result['entities']}

Summary:
{iep_result['summary']}
"""
            
            # Create new document with enhanced content
            results.append(LangchainDocument(
                page_content=enhanced_content,
                metadata={
                    **doc.metadata,
                    "processed_with": "dspy",
                    "has_enhanced_content": True
                }
            ))
            print(f"Successfully processed document with DSPy: {doc.metadata.get('source')}")
            
        except Exception as e:
            print(f"Error processing document with DSPy: {e}")
            # Fall back to original document if processing fails
            results.append(doc)

    return results

def clean_text(text):
    """Clean extracted text by removing common OCR artifacts and formatting issues."""
    # Remove non-standard characters and formatting artifacts
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
    text = re.sub(r'[.:;,]{2,}', '.', text)     # Clean up repeated punctuation
    text = re.sub(r'\s+', ' ', text)            # Normalize whitespace
    text = re.sub(r'\\n+', '\n', text)          # Normalize newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)     # Normalize paragraph breaks
    
    # Remove common OCR artifacts
    text = re.sub(r'[~`]{2,}', '', text)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', '', text)
    
    return text.strip()

def load_documents(data_path: str, use_dspy: bool = False) -> List[LangchainDocument]:
    """Load documents and optionally process them with DSPy.
    
    Args:
        data_path (str): Path to directory containing documents
        use_dspy (bool): Whether to process documents with DSPy pipeline
        
    Returns:
        List[LangchainDocument]: List of processed documents
    """
    docs = []
    print(f"Scanning directory: {data_path}")
    
    if not os.path.exists(data_path):
        print(f"Directory does not exist: {data_path}")
        return docs

    # Load documents as before
    for file in os.listdir(data_path):
        file_path = os.path.join(data_path, file)
        print(f"Processing file: {file_path}")
        
        try:
            if file.endswith((".docx", ".doc")):
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                text = clean_text(text)
                print(f"Extracted {len(text)} characters from DOCX")
                if text.strip():  # Only add if there's actual content
                    docs.append(LangchainDocument(
                        page_content=text, 
                        metadata={
                            "source": file,
                            "type": "docx",
                            "path": file_path
                        }
                    ))
            
            elif file.endswith(".pdf"):
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text = clean_text(text)
                        text_parts.append(text)
                
                text = "\n\n".join(text_parts)
                print(f"Extracted {len(text)} characters from PDF")
                if text.strip():  # Only add if there's actual content
                    docs.append(LangchainDocument(
                        page_content=text, 
                        metadata={
                            "source": file,
                            "type": "pdf",
                            "path": file_path,
                            "pages": len(reader.pages)
                        }
                    ))
            
            elif file.endswith((".txt", ".md")):
                with open(file_path, "r", encoding='utf-8') as file_content:
                    text = file_content.read()
                    text = clean_text(text)
                    print(f"Extracted {len(text)} characters from {file.split('.')[-1].upper()}")
                    if text.strip():  # Only add if there's actual content
                        docs.append(LangchainDocument(
                            page_content=text, 
                            metadata={
                                "source": file,
                                "type": "text",
                                "path": file_path
                            }
                        ))
            
        except Exception as e:
            print(f"Error processing file {file}: {str(e)}")
            continue
    
    print(f"Total documents loaded: {len(docs)}")
    
    # Process with DSPy if requested
    if use_dspy and docs:
        print("Processing documents with DSPy pipeline...")
        docs = process_with_dspy(docs)
        print(f"DSPy processing complete. Processed {len(docs)} documents.")
    
    if docs:
        print("Sample content from first document:")
        print(docs[0].page_content[:200] + "...")
    
    # Verify loaded documents
    for doc in docs:
        if not doc.page_content or not isinstance(doc.page_content, str):
            logger.warning(f"Invalid document content in {doc.metadata.get('source', 'unknown')}")
            docs.remove(doc)
            
    return docs